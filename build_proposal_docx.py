"""开题报告 md → docx (复用 build_docx 的样式)"""
import re, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:/Users/da983/Desktop/GraphVisualizer"
SRC  = os.path.join(ROOT, "开题报告.md")
DEST = os.path.join(ROOT, "开题报告.docx")

def set_run_font(run, name_zh="宋体", name_en="Times New Roman",
                 size_pt=12, bold=False):
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name_zh)
    rfonts.set(qn('w:ascii'), name_en)
    rfonts.set(qn('w:hAnsi'), name_en)

def set_pf(p, indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
           space_before=0, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_chars > 0:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.5)

# 大标题
p = doc.add_paragraph()
set_pf(p, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
r = p.add_run("本科生毕业论文(设计)开题报告")
set_run_font(r, name_zh="黑体", size_pt=18, bold=True)

with open(SRC, encoding='utf-8') as f:
    md = f.read()
lines = md.split('\n')

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    if not line.strip() or line.strip() == '---' or line.startswith('# '):
        i += 1; continue

    if line.startswith('## '):
        p = doc.add_paragraph()
        set_pf(p, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=14, space_after=8)
        r = p.add_run(line[3:].strip())
        set_run_font(r, name_zh="黑体", size_pt=14, bold=True)
        i += 1; continue

    if line.startswith('### '):
        p = doc.add_paragraph()
        set_pf(p, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=10, space_after=6)
        r = p.add_run(line[4:].strip())
        set_run_font(r, name_zh="黑体", size_pt=12, bold=True)
        i += 1; continue

    if line.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(cells)
            i += 1
        if len(rows) >= 2 and all(re.fullmatch(r':?-+:?', c) for c in rows[1]):
            rows.pop(1)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = 'Light Grid Accent 1'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci, txt in enumerate(row):
                    if ci >= len(rows[0]): continue
                    cell = t.cell(ri, ci); cell.text = ''
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(txt.replace('**', ''))
                    set_run_font(cr, size_pt=10.5, bold=(ri == 0 or '**' in txt))
        continue

    # 普通段
    p = doc.add_paragraph()
    set_pf(p, indent_chars=2)
    text = line
    parts = re.split(r'(\*\*.*?\*\*|\[\d+\])', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_run_font(r, size_pt=12, bold=True)
        elif re.fullmatch(r'\[\d+\]', part):
            r = p.add_run(part); r.font.superscript = True
            set_run_font(r, size_pt=12)
        else:
            r = p.add_run(part); set_run_font(r, size_pt=12)
    i += 1

doc.save(DEST)
print(f"已生成: {DEST}")
