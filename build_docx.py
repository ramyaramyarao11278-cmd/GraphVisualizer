"""
将 论文_重构稿.md 转换为符合学校模板的完整 docx
- 中文宋体小四 / 行距 1.5 / 段首缩进 2 字符
- 章标题黑体三号居中 / 节标题黑体小三 / 子节黑体四号
- 自动嵌入图片(从 figure_manifest.json 读映射)
- 表格自动居中
- 引用 [N] 自动设上标
- 页眉"山东师范大学本科生毕业论文"
"""
import re, json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:/Users/da983/Desktop/GraphVisualizer"
SRC  = os.path.join(ROOT, "论文_重构稿.md")
DEST = os.path.join(ROOT, "基于Qt的图算法可视化系统设计与实现_重构.docx")
MANIFEST = os.path.join(ROOT, "figure_manifest.json")

# 加载图注 → 图片映射
caption_to_image = {}
if os.path.exists(MANIFEST):
    with open(MANIFEST, encoding='utf-8') as f:
        caption_to_image = json.load(f)
print(f"加载了 {len(caption_to_image)} 项图片映射")

# ============== 字体辅助 ==============
def set_run_font(run, name_zh="宋体", name_en="Times New Roman",
                 size_pt=12, bold=False):
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name_zh)
    rfonts.set(qn('w:ascii'), name_en)
    rfonts.set(qn('w:hAnsi'), name_en)

def set_paragraph_format(p, indent_chars=2, line_spacing=1.5,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=0, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_chars > 0:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))

def add_run_with_supref(p, text):
    """处理段落中的 [N] 引用,自动设上标"""
    parts = re.split(r'(\[\d+\])', text)
    for part in parts:
        if not part: continue
        run = p.add_run(part)
        if re.fullmatch(r'\[\d+\]', part):
            run.font.superscript = True
            set_run_font(run, size_pt=12)
        else:
            set_run_font(run, size_pt=12)

# ============== 文档级设置 ==============
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("山东师范大学本科生毕业论文")
set_run_font(hr, size_pt=10.5)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============== 内容写入 ==============
def add_chapter_title(text):
    p = doc.add_paragraph()
    set_paragraph_format(p, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=24, space_after=18)
    r = p.add_run(text)
    set_run_font(r, name_zh="黑体", size_pt=16, bold=True)

def add_section_title(text, level=2):
    sizes = {2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    set_paragraph_format(p, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=12, space_after=6)
    r = p.add_run(text)
    set_run_font(r, name_zh="黑体", size_pt=sizes.get(level, 12), bold=True)

def add_body_paragraph(text):
    p = doc.add_paragraph()
    set_paragraph_format(p, indent_chars=2)
    add_run_with_supref(p, text)

def add_figure(caption):
    img_rel = caption_to_image.get(caption)
    img_abs = os.path.join(ROOT, img_rel) if img_rel else None
    # 图片
    p1 = doc.add_paragraph()
    set_paragraph_format(p1, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=6, space_after=2)
    if img_abs and os.path.exists(img_abs):
        run = p1.add_run()
        run.add_picture(img_abs, width=Inches(5.5))
    else:
        r = p1.add_run("[ 图片缺失: " + (img_rel or "未在 manifest 中") + " ]")
        set_run_font(r, size_pt=10.5)
        r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    # 标题
    p2 = doc.add_paragraph()
    set_paragraph_format(p2, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=2, space_after=12)
    r2 = p2.add_run(caption)
    set_run_font(r2, size_pt=10.5, bold=True)

def add_table_from_md(rows):
    if len(rows) < 2: return
    if all(re.fullmatch(r':?-+:?', c) for c in rows[1]):
        rows.pop(1)
    if not rows: return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= cols: continue
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(cell_text)
            set_run_font(r, size_pt=10.5, bold=(i == 0))

# ============== 解析 markdown ==============
with open(SRC, encoding='utf-8') as f:
    md = f.read()

md = re.split(r'\n---\n', md, maxsplit=1)[1]
lines = md.split('\n')

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line.strip() or line.strip() == '---':
        i += 1; continue
    if line.startswith('>'):
        i += 1; continue

    if line.startswith('# '):
        add_chapter_title(line[2:].strip())
        i += 1; continue

    if line.startswith('## '):
        title = line[3:].strip()
        if re.match(r'^(摘\s*要|Abstract|参考文献|致\s*谢)$', title):
            add_chapter_title(title)
        else:
            add_section_title(title, level=2)
        i += 1; continue

    if line.startswith('### '):
        add_section_title(line[4:].strip(), level=3)
        i += 1; continue

    if line.startswith('**关键词**') or line.startswith('**Keywords**'):
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        p = doc.add_paragraph()
        set_paragraph_format(p, indent_chars=2)
        r = p.add_run(text)
        set_run_font(r, size_pt=12, bold=True)
        i += 1; continue

    if line.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(cells)
            i += 1
        add_table_from_md(rows)
        continue

    m = re.match(r'^(图\s*\d+[-‒]\d+\s.*)$', line)
    if m:
        add_figure(m.group(1))
        i += 1; continue

    add_body_paragraph(line)
    i += 1

doc.save(DEST)

# 统计
import zipfile
with zipfile.ZipFile(DEST) as z:
    media = [n for n in z.namelist() if n.startswith('word/media/')]
print(f"\n生成: {DEST}")
print(f"段落: {len(doc.paragraphs)}  表格: {len(doc.tables)}  内嵌图片: {len(media)}")
